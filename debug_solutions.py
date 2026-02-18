import docx
import re

def find_solution_tables(file_path):
    doc = docx.Document(file_path)
    
    print("--- Buscando Patrones de Soluciones en TABLAS ---")
    
    # Regex para respuesta: "1", "1.", "1-", "a", "a)", "a.", etc.
    re_num = re.compile(r'^\d+[\.\-\)]?$')
    re_let = re.compile(r'^[abcdABCD][\.\-\)]?$')
    re_pair = re.compile(r'^\d+[\.\-\)]?\s*[abcdABCD]$') # "1a", "1 a"

    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        
        score = 0
        total_cells = 0
        
        sample_data = []

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                text = cell.text.strip()
                row_text.append(text)
                total_cells += 1
                
                if re_num.match(text) or re_let.match(text) or re_pair.match(text):
                    score += 1
            
            if len(sample_data) < 5: # Guardar primeros 5 para preview
                sample_data.append(row_text)
        
        if total_cells > 0:
            density = score / total_cells
            # Si más del 20% de las celdas parecen respuestas, es candidata
            if density > 0.2 and rows > 2:
                print(f"\n✅ TABLA CANDIDATA {i} (Densidad: {density:.2f})")
                print(f"Dimensiones: {rows}x{cols}")
                for r in sample_data:
                    print(f"  {r}")
                print("  ...")

    print("\n--- Buscando Patrones de Soluciones en TEXTO ---")
    # Buscar bloques de texto tipo "1-a 2-b 3-c"
    regex_bulk = re.compile(r'(\d+[\.\-\)]?\s*[abcdABCD])\s+(\d+[\.\-\)]?\s*[abcdABCD])')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        matches = regex_bulk.findall(text)
        if len(matches) > 3: # Si hay más de 3 pares en un párrafo
            print(f"\n✅ PÁRRAFO CANDIDATO {i}:")
            print(f"  {text[:200]}...")

if __name__ == "__main__":
    find_solution_tables('pinche test mad CM.docx')
